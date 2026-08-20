import base64
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
NOTEBOOK = ROOT / "Day2/실습자료/실습_LoRA_SFT_파인튜닝.ipynb"
CHART = ROOT / "Day2/실습자료/epoch_rank8_실험_결과.png"
OUTPUT = ROOT / "Day2/실습자료/LoRA_SFT_epoch_rank8_실행시간_결과보고서.docx"

INK = "243447"
BLUE = "2F5D8C"
LIGHT = "EAF1F8"
GRID = "CAD4DF"
MUTED = "66717E"
FONT = "Apple SD Gothic Neo"


def set_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_borders(table, color=GRID, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(sum(widths_dxa)))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table_borders(table)
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), font_size, True, BLUE)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), 15 if level == 1 else 12, True, BLUE if level == 1 else INK)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(text), 10)


def extract_last_png():
    notebook = json.loads(NOTEBOOK.read_text(encoding="utf-8"))
    images = [
        output["data"]["image/png"]
        for cell in notebook["cells"]
        for output in cell.get("outputs", [])
        if output.get("data", {}).get("image/png")
    ]
    if not images:
        raise RuntimeError("노트북에서 PNG 출력 이미지를 찾지 못했습니다.")
    raw = images[-1]
    if isinstance(raw, list):
        raw = "".join(raw)
    CHART.write_bytes(base64.b64decode(raw))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("SKALA LoRA SFT 실습  |  "), 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def build():
    extract_last_png()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    for name, before, after in (("Heading 1", 8, 5), ("Heading 2", 6, 3)):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    set_run(title.add_run("LoRA SFT 파인튜닝 결과 해석 보고서"), 22, True, INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    set_run(subtitle.add_run("Epoch 변화 및 LoRA Rank 16→8 축소 실험"), 12, True, BLUE)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    set_run(meta.add_run("Qwen2.5-0.5B-Instruct  |  Train 400건  |  Test 20건  |  Rank 8·16 비교  |  Seed 42"), 9, color=MUTED)

    add_heading(doc, "1. 학습 결과 요약")
    add_body(doc, "Baseline 규칙 준수율은 32.2%였고, rank 8·1 epoch LoRA SFT 이후 92.7%로 60.5%p 상승했다. 정답 SQL(completion)에만 loss를 적용했으며 train loss는 0.1376, eval loss는 0.0122였다.")
    add_table(doc, ["구분", "설정", "Train loss", "Eval loss", "규칙 준수율"], [
        ["Baseline", "학습 전", "-", "-", "32.2%"],
        ["Fine-Tuned", "400건 / 1 epoch / r=8", "0.1376", "0.0122", "92.7%"],
    ], [1750, 3150, 1350, 1350, 1760])

    add_heading(doc, "규칙별 변화")
    add_table(doc, ["규칙", "Baseline", "Fine-Tuned", "해석"], [
        ["alias / formula / orderdesc", "0%", "100%", "실매출 계산·별칭·정렬 패턴 학습"],
        ["completed", "10%", "100%", "완료 주문 필터가 안정화"],
        ["semicolon", "40%", "100%", "출력 형식 규칙 개선"],
        ["groupby", "0%", "75%", "대부분 개선, 차원 선택 오류 잔존"],
        ["period", "0%", "55%", "개선됐으나 가장 낮은 준수율"],
        ["select_only / table / limit", "100%", "100%", "Baseline부터 이미 충족"],
    ], [2200, 1250, 1450, 4560])

    add_body(doc, "관찰 사례: ‘7월 지역 기준’ 질문에서 Fine-Tuned 모델은 계산식, 완료 상태, 기간, 정렬은 맞췄지만 GROUP BY category를 출력했다. 즉 SQL 골격은 학습했으나 자연어의 집계 차원(region/category)을 정확히 연결하는 문제는 남았다.", "관찰 사례:")

    doc.add_page_break()
    add_heading(doc, "2. Epoch·LoRA Rank·실행 시간 결과")
    add_body(doc, "학습 데이터 400건과 seed 42를 고정하고 rank를 16에서 8로 줄였다. Rank 8의 학습 가능 파라미터는 540,672개(0.1093%)로 rank 16의 1,081,344개(0.2184%) 대비 정확히 절반이다.")
    add_table(doc, ["Epoch", "Rank", "Train loss", "Eval loss", "준수율", "실행 시간"], [
        ["1", "8", "0.1376", "0.0122", "92.7%", "244.32초"],
        ["3", "8", "0.0423", "0.0047", "97.8%", "210.29초"],
        ["10", "8", "0.0125", "0.0006", "100.0%", "316.20초"],
    ], [900, 900, 1500, 1450, 1450, 3160])

    add_table(doc, ["Epoch", "Rank 16 준수율", "Rank 8 준수율", "차이", "해석"], [
        ["1", "98.3%", "92.7%", "-5.6%p", "초기 적응은 rank 16 우세"],
        ["3", "100.0%", "97.8%", "-2.2%p", "반복 학습으로 격차 축소"],
        ["10", "100.0%", "100.0%", "0.0%p", "낮은 rank도 최종 성능 도달"],
    ], [900, 1700, 1700, 1300, 3760], font_size=8.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    picture = p.add_run().add_picture(str(CHART), width=Inches(5.2))
    doc_pr = picture._inline.docPr
    doc_pr.set("title", "Rank 8 Epoch별 LoRA SFT 실험 결과")
    doc_pr.set("descr", "Rank 8에서 Epoch 1, 3, 10에 따른 train loss, eval loss, 규칙 준수율 및 실행 시간 그래프")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    set_run(cap.add_run("그림 1. Rank 8의 Epoch별 loss·준수율·실행 시간 (노트북 마지막 이미지)"), 8.5, color=MUTED)

    add_heading(doc, "개인 해석")
    add_body(doc, "Rank 8은 1·3 epoch에서 rank 16보다 각각 5.6%p, 2.2%p 낮았지만 10 epoch에는 동일한 100%에 도달했다. 즉 rank 축소는 초기 학습 속도를 늦췄으나 충분히 반복하면 이번 과제 수준의 규칙을 표현할 용량은 확보했다. 파라미터 효율을 우선하면 rank 8·10 epoch, 빠른 성능 도달을 우선하면 rank 16·3 epoch가 후보가 된다.")
    add_body(doc, "실행 시간은 1→3 epoch에서 오히려 244.32→210.29초로 줄고 10 epoch에서 316.20초로 늘었다. 모델 로딩, 평가, 20건 SQL 생성, GPU 상태가 모두 포함된 1회 실측이므로 epoch와 순수 학습 시간의 선형 관계로 해석할 수 없다. 반복 측정의 중앙값과 trainer의 train_runtime을 별도로 기록해야 공정하다.")

    doc.add_page_break()
    add_heading(doc, "3. 조별 토의 및 고찰 (취합용 초안)")
    add_body(doc, "A(epoch) 결과와 추가 rank 16→8 비교를 함께 보면, 작은 rank는 적은 epoch에서 불리하지만 반복 학습으로 격차를 회복했다. 따라서 성능만이 아니라 학습 파라미터 수와 목표 성능 도달 시간까지 함께 비교해야 한다.")

    add_heading(doc, "다른 실험과 비교해 새롭게 확인할 점", 2)
    add_bullet(doc, "B(데이터 크기): 50/100/400건에서 소규모 데이터가 period·groupby처럼 문맥 의존 규칙을 얼마나 놓치는지 비교한다. 400건보다 낮은 준수율이 크다면 데이터 다양성이 반복 학습보다 중요하다는 근거가 된다.")
    add_bullet(doc, "C(LoRA rank): 이번 r=8/16 비교에서는 높은 rank가 초기 수렴에 유리했지만 10 epoch에서는 차이가 사라졌다. r=2/32까지 추가하면 최소 충분 용량과 과도한 용량을 확인할 수 있다.")
    add_bullet(doc, "공정 비교: train_size·epoch·rank 중 하나만 바꾼 결과끼리 비교하고, 동일 seed와 동일 평가 문항을 유지한다.")

    add_heading(doc, "토의 후 보완할 의견", 2)
    add_body(doc, "초기에는 rank를 낮추면 최종 성능도 떨어질 것으로 예상했지만, rank 8은 10 epoch에서 100%에 도달했다. 이에 따라 rank는 절대적인 최종 성능보다 ‘얼마나 빨리 목표 성능에 도달하는가’에 더 큰 영향을 줄 수 있다는 쪽으로 생각을 보완했다.")

    add_heading(doc, "조의 결론 초안: 가장 큰 영향 요인", 2)
    add_body(doc, "현재 결론은 충분하고 다양한 학습 데이터가 가장 중요한 요인이며, 그다음은 epoch와 rank의 조합이다. Rank가 작으면 더 많은 반복이 필요하고, rank가 크면 적은 epoch에서도 빠르게 높은 준수율에 도달했다. 다만 100% 도달 후의 loss 감소는 추가 규칙 성능으로 이어지지 않았으므로 ‘최소 비용으로 목표 준수율을 달성하는 조합’을 선택하는 것이 합리적이다.")

    add_table(doc, ["실험", "취합할 핵심 수치", "A 실험과의 비교 질문"], [
        ["B: 데이터 크기", "50/100/400건 준수율·loss", "적은 데이터에서 어떤 규칙이 먼저 무너지는가?"],
        ["C: LoRA rank", "r=2/8/16/32 준수율·시간", "목표 성능에 도달하는 최소 rank는 무엇인가?"],
        ["최종 결론", "최고 성능·파라미터·시간", "성능과 비용의 균형이 가장 좋은 조합은 무엇인가?"],
    ], [1800, 2850, 4710], font_size=8.9)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(0)
    set_run(note.add_run("해석 한계  "), 9, True, BLUE)
    set_run(note.add_run("규칙 준수율은 SQL 실행 정확도나 의미적 정답률과 동일하지 않으며, 표본 20건의 자동 규칙 검사 결과다. 향후에는 별도 hold-out과 SQL 실행 결과 일치 여부를 함께 평가하는 것이 바람직하다."), 9, color=MUTED)

    doc.core_properties.title = "LoRA SFT 결과 보고서 - Epoch 및 Rank 16→8 비교"
    doc.core_properties.subject = "Baseline, Epoch, LoRA Rank 및 실행 시간 분석"
    doc.core_properties.author = "SKALA 실습"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

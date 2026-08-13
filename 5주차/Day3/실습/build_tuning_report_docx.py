from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_tuning_report import sections


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "SQL_쿼리_튜닝_실행결과_입력용.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
MUTED = "666666"
CODE_FILL = "F4F6F9"
RESULT_FILL = "E8EEF5"
POINT_FILL = "F1F7E9"

NON_INDEX = {
    3: {
        "title": "조건식 재작성: split_part()",
        "queries": """EXPLAIN (ANALYZE, BUFFERS)
SELECT * FROM hr.employees
WHERE split_part(email, '@', 2) = 'gmail.com';

EXPLAIN (ANALYZE, BUFFERS)
SELECT employee_id, first_name, last_name, email
FROM hr.employees
WHERE split_part(email, '@', 2) = 'outlook.com';""",
        "results": """gmail: LIKE 2.206 ms -> split_part() 2.363 ms
outlook: LIKE 2.119 ms -> split_part() 2.545 ms
Scan: Seq Scan 유지 / Buffers: shared hit=786

판단: 행별 함수 계산은 추가되었지만 전체 스캔이 제거되지 않아 채택하지 않음.""",
        "point": "접미사 검색은 쿼리 재작성만으로 개선되지 않았다. 반복 검색이라면 도메인 컬럼 정규화 또는 reverse(email) 표현식 인덱스를 검토한다.",
    },
    4: {
        "title": "SELECT * 제거와 조회 컬럼 축소",
        "queries": """EXPLAIN (ANALYZE, BUFFERS)
SELECT employee_id, email, hire_date, salary
FROM hr.employees
WHERE hire_date >= CURRENT_DATE - INTERVAL '365 days'
  AND status = 'ACTIVE'
ORDER BY salary DESC
LIMIT 100;""",
        "results": """광범위 컬럼: 5.304 ms / top-N Memory 47 kB
필요 컬럼: 4.768 ms / top-N Memory 36 kB
실행 시간: 약 10.1% 개선
정렬 메모리: 11 kB 감소

판단: 행 폭과 전송량은 줄었지만 Seq Scan과 top-N Sort는 유지됨.""",
        "point": "화면에 필요한 컬럼만 조회하는 것은 유효한 비인덱스 튜닝이다. 다만 핵심 병목인 스캔과 정렬을 없애지는 못한다.",
    },
    5: {
        "title": "IN 단순화와 상호 배타 UNION ALL",
        "queries": """EXPLAIN (ANALYZE, BUFFERS)
SELECT * FROM hr.employees
WHERE department_id = 10 OR job_id BETWEEN 3 AND 5;

EXPLAIN (ANALYZE, BUFFERS)
SELECT * FROM hr.employees WHERE department_id = 10
UNION ALL
SELECT * FROM hr.employees
WHERE job_id IN (3,4,5) AND department_id <> 10;""",
        "results": """인덱스 없음: IN 5.658 ms -> BETWEEN 3.938 ms
  Seq Scan과 Rows Removed by Filter 45,915건은 유지
인덱스 있음: OR BitmapOr 0.412 ms -> UNION ALL 0.563 ms

판단: BETWEEN은 반복 측정이 필요하며, UNION ALL은 더 느려 채택하지 않음.""",
        "point": "조건식 재작성은 결과 집합이 동일한지 먼저 확인해야 한다. 현재 데이터에서는 원본 OR + BitmapOr가 가장 안정적이다.",
    },
}


def font(run, name="Calibri", size=11, color="000000", bold=False):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd")) or OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    if node.getparent() is None:
        properties.append(node)


def margins(cell, value=120):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:tcMar")) or OxmlElement("w:tcMar")
    if node.getparent() is None:
        properties.append(node)
    for side in ("top", "start", "bottom", "end"):
        item = node.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        if item.getparent() is None:
            node.append(item)


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    table_node = table._tbl
    properties = table_node.tblPr
    width_node = properties.find(qn("w:tblW")) or OxmlElement("w:tblW")
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")
    if width_node.getparent() is None:
        properties.append(width_node)
    indent_node = properties.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    indent_node.set(qn("w:w"), "120")
    indent_node.set(qn("w:type"), "dxa")
    if indent_node.getparent() is None:
        properties.append(indent_node)
    grid = table_node.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:w"), str(widths[index]))
            cell._tc.get_or_add_tcPr().tcW.set(qn("w:type"), "dxa")


def block(doc, text, fill, size=8.2):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.rstrip().splitlines()):
        if index:
            paragraph.add_run().add_break()
        font(paragraph.add_run(line), "Menlo", size, "1F2937")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def comparison_table(doc, rows):
    table = doc.add_table(rows=1, cols=6)
    set_table_geometry(table, [1800, 1620, 1260, 1620, 1260, 1800])
    headers = ("테스트", "튜닝 전", "전 시간", "튜닝 후", "후 시간", "개선율")
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, RESULT_FILL)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        font(cell.paragraphs[0].add_run(value), size=8.5, color=DARK, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if index >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            font(paragraph.add_run(value), size=8.2, color="08783E" if index == 5 else "1F2937", bold=index == 5)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 26, DARK, 0, 8),
        ("Heading 1", 16, BLUE, 12, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 11, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    font(header.add_run("SKALA | PostgreSQL 쿼리 튜닝 실습"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("2026-08-13"), size=9, color=MUTED)


def add_test(doc, phase, label, query, result):
    doc.add_heading(f"{phase} 테스트 {label}", level=3)
    doc.add_paragraph("쿼리문")
    block(doc, query, CODE_FILL, 8.0)
    doc.add_paragraph("출력결과")
    block(doc, result, RESULT_FILL, 8.0)


def result_placeholder(phase, label):
    return (
        f"[{phase} 테스트 {label} 실행 결과를 여기에 붙여넣기]\n"
        "\n\n\n\n\n\n"
        "Planning Time: ____________________ ms\n"
        "Execution Time: ___________________ ms"
    )


def add_non_index_test(doc, number):
    item = NON_INDEX.get(number)
    if not item:
        return
    doc.add_heading("8. 인덱스 이외의 튜닝 실제 적용", level=2)
    doc.add_heading(item["title"], level=3)
    doc.add_paragraph("실행 쿼리")
    block(doc, item["queries"], CODE_FILL, 8.0)
    doc.add_paragraph("실측 결과 및 판단")
    block(doc, item["results"], RESULT_FILL, 8.2)
    doc.add_paragraph("최종 튜닝 Point")
    block(doc, item["point"], POINT_FILL, 9.2)


def build():
    doc = Document()
    configure(doc)
    doc.add_heading("SQL 쿼리 튜닝 실습", 0)
    subtitle = doc.add_paragraph("튜닝 전·후 실행 계획 및 성능 비교 보고서")
    font(subtitle.add_run(), size=14, color=MUTED)
    meta = doc.add_paragraph()
    font(meta.add_run("Database: skala_db | Schema: hr | PostgreSQL 17.10\n"), size=10, color=MUTED)
    font(meta.add_run("구성: 문제 · 튜닝 전 쿼리/결과 · 튜닝 내역 · 튜닝 후 쿼리/결과 · 의견"), size=10, color=MUTED)
    doc.add_paragraph("2번의 수기 작성 흐름을 기준으로 3~5번의 조건, 튜닝 전 쿼리, 튜닝 방법, 튜닝 후 쿼리, 성능, 추가 포인트, 결론 순서를 통일했습니다.")
    doc.add_paragraph("제출 전 입력: 작성자 ____________________    반 ________")

    for section in sections:
        doc.add_heading(f'{section["no"]}. {section["title"]}', level=1)
        doc.add_heading("1. 문제", level=2)
        doc.add_paragraph(section["goal"])

        doc.add_heading("2. 튜닝 전 쿼리 및 출력결과", level=2)
        for label, query, result in section["before"]:
            add_test(doc, "튜닝 전", label, query, result if section["no"] == 2 else result_placeholder("튜닝 전", label))

        doc.add_heading("3. 튜닝 작업 내역", level=2)
        block(doc, section["ddl"], CODE_FILL, 8.0)

        doc.add_heading("4. 튜닝 후 쿼리 및 출력결과", level=2)
        for label, query, result in section["after"]:
            add_test(doc, "튜닝 후", label, query, result if section["no"] == 2 else result_placeholder("튜닝 후", label))

        doc.add_heading("5. 성능 비교", level=2)
        if section["no"] == 2:
            comparison_table(doc, section["comparison"])
        else:
            comparison_table(doc, [(label, "____________", "______ ms", "____________", "______ ms", "______ %") for label, *_ in section["comparison"]])
        doc.add_heading("6. 추가 포인트 및 조별 의견", level=2)
        doc.add_paragraph(section["discussion"] if section["no"] == 2 else "[실행 계획, 스캔 방식, 버퍼, 제거된 행, 선택도를 비교한 조별 의견을 작성하세요.]\n\n\n")
        doc.add_heading("7. 결론 및 최적의 튜닝 Point", level=2)
        block(doc, section["best"] if section["no"] == 2 else "[튜닝 전·후 실행 결과를 근거로 결론과 최적의 튜닝 Point를 작성하세요.]\n\n\n", POINT_FILL, 9.5)
        add_non_index_test(doc, section["no"])

    doc.add_heading("종합 결론", level=1)
    for section in sections:
        paragraph = doc.add_paragraph()
        font(paragraph.add_run(f'{section["no"]}번. '), color=DARK, bold=True)
        font(paragraph.add_run(section["best"]))
    doc.add_paragraph("실행 시간은 캐시와 동시 부하에 따라 달라질 수 있으므로, 실행 시간과 함께 Seq Scan, Sort, Rows Removed by Filter, Buffers 변화를 판단해야 합니다.")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

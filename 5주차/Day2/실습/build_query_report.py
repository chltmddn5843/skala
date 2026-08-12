import ast
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "종합실습2_queries.py"
LOG = max((ROOT / "logs").glob("종합실습2_all_*.log"), key=lambda p: p.stat().st_mtime)
OUTPUT = ROOT / "종합실습2_쿼리_실행결과_보고서.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
MUTED = "666666"
CODE_FILL = "F4F6F9"
RESULT_FILL = "E8EEF5"


def font(run, name="Calibri", size=11, color="000000", bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd")) or OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    if node.getparent() is None:
        properties.append(node)


def margins(cell, value=120):
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:tcMar")) or OxmlElement("w:tcMar")
    if node.getparent() is None:
        properties.append(node)
    for side in ("top", "start", "bottom", "end"):
        item = node.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        if item.getparent() is None:
            node.append(item)


def block(doc, text, fill, size=8.5):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(text.rstrip().splitlines()):
        if index:
            paragraph.add_run().add_break()
        font(paragraph.add_run(line), "Menlo", size, "1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def queries_and_problems():
    tree = ast.parse(SOURCE.read_text(encoding="utf-8"))
    result = {}
    for node in tree.body:
        if not isinstance(node, ast.FunctionDef) or not re.fullmatch(r"q\d{2}_.+", node.name):
            continue
        number = str(int(node.name[1:3]))
        problem = ast.get_docstring(node).split(". ", 1)[-1].rstrip(".")
        queries = []
        for child in ast.walk(node):
            if isinstance(child, ast.Constant) and isinstance(child.value, str):
                sql = child.value.strip()
                if re.search(r"\b(SELECT|CREATE TABLE)\b", sql, re.I):
                    queries.append(sql)
        result[number] = (problem, queries)
    return result


def log_results():
    text = LOG.read_text(encoding="utf-8")
    matches = list(re.finditer(r"(?m)^\[(\d+)(?:-\d+)?\. ([^]]+)\]\n", text))
    result = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else text.find("\n로그 저장:", match.end())
        number, title = match.group(1), match.group(2)
        output = text[match.end():end].strip()
        if number in result:
            first_title, previous = result[number]
            result[number] = (first_title, f"{previous}\n\n[{title}]\n{output}")
        else:
            result[number] = (title, output)
    return result


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 26, DARK, 0, 8),
        ("Heading 1", 16, BLUE, 12, 8),
        ("Heading 2", 13, BLUE, 10, 6),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    font(header.add_run("SKALA | SQL 종합실습 2"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("2026-08-12"), size=9, color=MUTED)


def build():
    items = queries_and_problems()
    results = log_results()
    assert set(items) == set(results), "쿼리와 로그 항목이 일치하지 않습니다."

    doc = Document()
    configure(doc)
    doc.add_heading("SQL 종합실습 2", 0)
    subtitle = doc.add_paragraph("쿼리문 및 실행 결과 보고서")
    font(subtitle.add_run(), size=14, color=MUTED)
    meta = doc.add_paragraph()
    font(meta.add_run(f"실행 로그: {LOG.name}\n"), size=10, color=MUTED)
    font(meta.add_run("구성: 문제 · 쿼리문 · 출력결과 | PostgreSQL / lab 스키마"), size=10, color=MUTED)
    doc.add_paragraph("본 보고서는 터미널 캡처 대신 실제 실행 로그를 사용하여 재현 가능한 형태로 정리했습니다.")

    for number in sorted(items, key=int):
        doc.add_page_break()
        problem, queries = items[number]
        title, output = results[number]
        doc.add_heading(f"{number}. {title}", level=1)
        doc.add_heading("1. 문제", level=2)
        doc.add_paragraph(problem)
        doc.add_heading("2. 쿼리문", level=2)
        for query in queries:
            block(doc, query, CODE_FILL, 8.0)
        doc.add_heading("3. 출력결과", level=2)
        block(doc, output, RESULT_FILL, 7.5 if number == "13" else 8.2)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

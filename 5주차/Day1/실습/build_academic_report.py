from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ERD = ROOT / "image" / "인사이트" / "1786435268484.png"
OUT = ROOT / "학사관리시스템_DB_설계_요약보고서.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.1, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, bold=False, color="000000", size=11, after=6, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, after=after, align=align)
    set_run_font(p.add_run(text), size=size, color=color, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.0)
        set_run_font(p.add_run(header), size=9.5, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, after=0, line=1.0)
            set_run_font(p.add_run(str(value)), size=9.3)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.1)
    set_run_font(p.add_run(f"{label}  "), size=10.5, color=DARK_BLUE, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.0)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        set_run_font(p.add_run(line), name="Consolas", size=8.2, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_capture_placeholder(doc, number, title, guidance):
    add_text(doc, f"캡처 {number}. {title}", bold=True, color=DARK_BLUE, size=10.5, after=3)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=12, after=12, line=1.1)
    set_run_font(p.add_run("[DBeaver 캡처 이미지를 여기에 삽입]"), size=12, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p2, after=0)
    set_run_font(p2.add_run(guidance), size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(header, after=0)
    set_run_font(header.add_run("학사관리시스템 데이터베이스 설계 요약보고서"), size=9, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def build_report():
    doc = Document()
    configure_document(doc)

    # 첫 페이지: memo_masthead
    add_text(doc, "DATABASE DESIGN REPORT", bold=True, color=BLUE, size=10, after=10)
    add_text(doc, "학사관리시스템 데이터베이스\n설계 및 실습 요약보고서", bold=True, color=INK, size=24, after=8)
    add_text(doc, "PostgreSQL · DBeaver · Bridge Model", color=GRAY, size=13, after=18)

    metadata = [
        ("주제", "강좌 개설·수강신청·성적 처리 중심 학사관리시스템"),
        ("DBMS", "PostgreSQL"),
        ("스키마", "academic"),
        ("핵심 테이블", "users, terms, courses, course_offerings, enrollments"),
        ("작성일", "2026년 8월 11일"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        style_paragraph(p, after=3)
        set_run_font(p.add_run(f"{label}: "), size=10.5, color=DARK_BLUE, bold=True)
        set_run_font(p.add_run(value), size=10.5)

    add_callout(
        doc,
        "핵심 요약",
        "과목의 고정 정보와 학기별 개설 정보를 분리하고, enrollments를 학생과 개설 강좌 사이의 Bridge 엔티티로 설계했다.",
    )

    add_heading(doc, "1. 설계 목적과 범위", 1)
    add_text(
        doc,
        "본 실습은 학사관리시스템 전체를 구현하는 대신 강좌 개설, 수강신청, 성적 처리라는 세 가지 핵심 업무에 집중한다. 기능 범위를 줄여 엔티티의 책임, 관계, 제약조건이 업무 시나리오와 직접 연결되도록 설계했다.",
    )
    for item in (
        "관리자: 학기·과목·교수를 선택해 분반을 개설하고 신청 상태를 관리한다.",
        "학생: 신청 가능한 강좌를 조회하고 정원을 확인한 뒤 수강신청한다.",
        "교수: 담당 강좌의 수강생에게 점수와 등급을 입력하고 공개한다.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "2. 요구사항과 업무 시나리오", 1)
    add_heading(doc, "2.1 강좌 개설", 2)
    add_text(doc, "관리자가 특정 학기에 담당 교수, 분반, 정원, 강의실을 지정해 강좌를 개설한다.")
    add_text(doc, "관리자 로그인 → 학기 선택 → 과목 선택 → 교수 선택 → 분반·정원·강의실 입력 → 강좌 생성 → OPEN", color=DARK_BLUE, size=10.5)

    add_heading(doc, "2.2 수강신청", 2)
    add_text(doc, "학생이 현재 학기의 OPEN 강좌를 조회하고 정원이 남은 분반에 신청한다.")
    add_text(doc, "학생 로그인 → 신청 가능 학기 확인 → 강좌 조회 → 정원 확인 → enrollments 저장", color=DARK_BLUE, size=10.5)

    add_heading(doc, "2.3 성적 처리", 2)
    add_text(doc, "교수가 본인의 담당 강좌 수강생에게 성적을 입력하고, 검토 후 학생에게 공개한다.")
    add_text(doc, "교수 로그인 → 담당 강좌 조회 → 수강생 조회 → 점수·등급 입력 → DRAFT → PUBLISHED", color=DARK_BLUE, size=10.5)

    add_heading(doc, "2.4 기능 범위", 2)
    add_table(
        doc,
        ["구분", "포함", "제외"],
        [
            ("사용자", "학생·교수·관리자 역할", "학과·전공·휴학 상태"),
            ("강좌", "학기·과목·교수·분반·정원·강의실", "공동교수·시간표 충돌"),
            ("수강", "신청·중복 방지·정원 확인", "수강취소·대기열·선수과목"),
            ("성적", "점수·등급·초안·공개", "평가항목별 점수·이의신청"),
        ],
        [1500, 3930, 3930],
    )

    doc.add_page_break()

    add_heading(doc, "3. ERD 설계 결과", 1)
    add_text(doc, "그림 1은 academic 스키마의 5개 엔티티와 외래키 관계를 나타낸다.", color=GRAY, size=9.5, after=4)
    if ERD.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=4)
        p.add_run().add_picture(str(ERD), width=Inches(6.35))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(caption, after=10)
        set_run_font(caption.add_run("그림 1. 학사관리시스템 ERD"), size=9.5, color=GRAY, italic=True)

    add_heading(doc, "3.1 ERD 범례", 2)
    add_table(
        doc,
        ["표기", "의미", "본 설계에서의 역할"],
        [
            ("PK", "기본키", "각 행의 안정적인 대표 식별자"),
            ("FK", "외래키", "사용자·학기·과목·개설 강좌 연결"),
            ("NN", "NOT NULL", "업무 수행에 필수인 값"),
            ("UK", "UNIQUE", "로그인 ID·과목 코드·중복 신청 방지"),
            ("1:N", "일대다 관계", "과목→개설강좌, 강좌→수강신청"),
        ],
        [1100, 2500, 5760],
    )

    add_heading(doc, "3.2 핵심 관계", 2)
    for item in (
        "courses 1:N course_offerings - 하나의 과목은 여러 학기와 분반으로 개설될 수 있다.",
        "terms 1:N course_offerings - 한 학기에는 여러 개설 강좌가 존재한다.",
        "users(PROFESSOR) 1:N course_offerings - 현재 범위에서는 강좌당 교수 한 명을 직접 참조한다.",
        "users(STUDENT) N:M course_offerings - enrollments가 다대다 관계를 해소한다.",
        "course_offerings 1:N enrollments - 한 강좌에 여러 학생이 신청한다.",
    ):
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "4. 엔티티 구성과 설계 근거", 1)
    add_table(
        doc,
        ["엔티티", "주요 속성", "책임과 설계 근거"],
        [
            ("users", "user_id, login_id, role", "공통 사용자와 권한 구분. role은 STUDENT·PROFESSOR·ADMIN만 허용"),
            ("terms", "year, semester, 신청 기간, status", "학기 중복을 막고 수강신청 가능 기간과 상태를 관리"),
            ("courses", "course_code, name, credits", "학기와 무관한 과목 고정 정보를 보관"),
            ("course_offerings", "course·term·professor FK, 분반·정원·강의실", "특정 학기의 실제 운영 단위. 동일 학기·과목·분반 중복 방지"),
            ("enrollments", "student·offering FK, 성적, 공개 상태", "학생-개설강좌 Bridge. 신청 관계와 수강 결과를 함께 관리"),
        ],
        [1600, 3000, 4760],
    )

    add_heading(doc, "4.1 주요 무결성 규칙", 2)
    add_table(
        doc,
        ["업무 규칙", "DB 제약조건", "목적"],
        [
            ("사용자 ID 중복 금지", "UNIQUE(login_id)", "로그인 식별 충돌 방지"),
            ("학기 중복 금지", "UNIQUE(year, semester)", "동일 학기 중복 생성 방지"),
            ("분반 중복 금지", "UNIQUE(term_id, course_id, section_no)", "같은 과목·학기·분반 중복 방지"),
            ("중복 수강신청 금지", "UNIQUE(student_id, offering_id)", "동일 학생의 동일 강좌 재신청 방지"),
            ("성적 범위", "CHECK(score BETWEEN 0 AND 100)", "유효하지 않은 점수 방지"),
            ("공개 성적 완전성", "PUBLISHED이면 점수·등급 NOT NULL", "불완전한 성적 공개 방지"),
        ],
        [2600, 3300, 3460],
    )
    add_callout(doc, "애플리케이션 책임", "교수·학생 역할 검증과 동시 수강신청의 정원 초과 방지는 단순 FK/CHECK만으로 해결할 수 없어 서비스 로직과 트랜잭션에서 처리한다.")

    doc.add_page_break()

    add_heading(doc, "5. 테이블 상호작용과 기능 플로우", 1)
    add_heading(doc, "5.1 강좌 개설 플로우", 2)
    for step in (
        "users에서 ADMIN 역할을 확인한다.",
        "terms에서 대상 학기를, courses에서 과목을 조회한다.",
        "users에서 PROFESSOR 역할 사용자를 선택한다.",
        "course_offerings에 분반·정원·강의실과 세 개의 FK를 저장한다.",
        "검토 후 status를 PLANNED에서 OPEN으로 변경한다.",
    ):
        add_number(doc, step)

    add_heading(doc, "5.2 수강신청 플로우", 2)
    for step in (
        "users에서 STUDENT 역할을 확인한다.",
        "terms의 기간과 상태를 기준으로 신청 가능한 학기를 찾는다.",
        "course_offerings와 courses를 JOIN해 OPEN 강좌를 보여준다.",
        "개설 강좌 행을 잠그고 enrollments 건수를 세어 잔여 정원을 확인한다.",
        "정원이 남으면 enrollments에 학생과 개설 강좌의 관계를 저장한다.",
    ):
        add_number(doc, step)

    add_code(doc, "BEGIN;\nSELECT capacity FROM academic.course_offerings\nWHERE offering_id = :offering_id FOR UPDATE;\n-- 현재 신청 인원 확인 후 INSERT\nCOMMIT;")

    add_heading(doc, "5.3 성적 처리 플로우", 2)
    for step in (
        "course_offerings.professor_id로 교수의 담당 강좌를 조회한다.",
        "enrollments와 users를 JOIN해 해당 강좌의 수강생을 조회한다.",
        "점수와 등급을 입력하고 grade_status를 DRAFT로 저장한다.",
        "검토가 끝나면 PUBLISHED로 변경한다.",
        "학생 조회 화면에서는 PUBLISHED 성적만 노출한다.",
    ):
        add_number(doc, step)

    doc.add_page_break()

    add_heading(doc, "6. PostgreSQL 구현 및 실습 결과", 1)
    add_text(doc, "구현은 두 개의 SQL 파일로 분리했다. DDL은 구조와 제약조건을 생성하고, 실습 쿼리는 샘플 데이터와 조회 결과를 만든다.")
    add_table(
        doc,
        ["파일", "역할", "주요 내용"],
        [
            ("academic_bridge_model.sql", "DDL", "스키마·5개 테이블·PK/FK/UNIQUE/CHECK·인덱스 생성"),
            ("academic_practice_queries.sql", "DML/조회", "테이블별 10건 이상 INSERT, 기본 조회, 함수, JOIN"),
            ("academic_bridge_model.dbml", "ERD", "dbdiagram.io에서 관계도 시각화"),
        ],
        [2600, 1700, 5060],
    )

    add_heading(doc, "6.1 샘플 데이터 기대 결과", 2)
    add_table(
        doc,
        ["테이블", "기대 건수", "구성"],
        [
            ("users", "14", "관리자 1, 교수 3, 학생 10"),
            ("terms", "10", "2022~2026년 봄·가을"),
            ("courses", "10", "컴퓨터·DB·AI 관련 과목"),
            ("course_offerings", "10", "2026년 가을 개설 분반"),
            ("enrollments", "10", "공개·초안·미입력 성적 포함"),
        ],
        [2600, 1600, 5160],
    )

    add_heading(doc, "6.2 실습 SQL 범위", 2)
    for item in (
        "SELECT + WHERE + ORDER BY: 신청 가능한 강좌를 과목 코드순으로 조회",
        "COALESCE: NULL 성적을 '미입력'으로 변환",
        "CASE WHEN: 성적 상태를 사용자용 한글 문구로 변환",
        "날짜 함수: 신청 시각을 한국 시간으로 표시하고 경과 일수 계산",
        "JOIN: enrollments를 중심으로 학생·강좌·과목·학기·교수를 교차 조회",
    ):
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "7. 제출용 캡처 체크리스트", 1)
    add_text(doc, "아래 영역은 수정 가능한 자리표시자다. DBeaver에서 해당 결과를 실행한 뒤 캡처 이미지를 선택해 교체하면 된다.")
    add_capture_placeholder(doc, 1, "PostgreSQL 접속 확인", "current_database(), current_user, version() 결과가 보이도록 캡처")
    add_capture_placeholder(doc, 2, "스키마와 5개 테이블 생성 결과", "DBeaver Navigator의 academic 스키마와 테이블 목록이 함께 보이도록 캡처")

    doc.add_page_break()
    add_heading(doc, "7. 제출용 캡처 체크리스트 (계속)", 1)
    add_capture_placeholder(doc, 3, "테이블별 10건 이상 INSERT 결과", "users 14건, 나머지 테이블 10건이 표시된 건수 검증 결과")
    add_capture_placeholder(doc, 4, "SELECT + WHERE + ORDER BY 결과", "2026년 가을 OPEN 강좌가 course_code 순으로 정렬된 결과")
    add_capture_placeholder(doc, 5, "COALESCE / CASE WHEN / 날짜 함수 결과", "미입력·작성 중·공개 완료 문구와 날짜 계산 결과")

    doc.add_page_break()
    add_heading(doc, "7. 제출용 캡처 체크리스트 (계속)", 1)
    add_capture_placeholder(doc, 6, "수강신청 Bridge JOIN 결과", "학생·학기·과목·교수·성적이 한 결과에 나타나는 JOIN 화면")

    add_heading(doc, "8. 결론", 1)
    add_text(
        doc,
        "본 설계는 기능 범위를 세 가지 업무로 제한해 테이블 책임과 관계를 명확히 했다. courses와 course_offerings를 분리해 과목과 학기별 운영 정보를 구분했고, enrollments를 Bridge 엔티티로 두어 학생과 강좌의 N:M 관계 및 성적 상태를 하나의 수강 맥락에서 관리했다.",
    )
    add_text(
        doc,
        "PK, FK, UNIQUE, CHECK를 통해 행 식별·참조 무결성·중복 방지·상태 유효성을 보장했다. 정원 초과와 역할 검증처럼 여러 행이나 업무 권한이 필요한 규칙은 애플리케이션 트랜잭션의 책임으로 분리했다.",
    )

    add_heading(doc, "부록. 제출 전 확인", 1)
    for item in (
        "ERD 관계선과 범례가 선명하게 보이는가?",
        "PostgreSQL 접속 결과 화면이 포함됐는가?",
        "각 SQL문과 실행 결과 화면이 한 쌍으로 제시됐는가?",
        "테이블별 최소 10건 이상의 데이터가 확인되는가?",
        "JOIN 결과에서 enrollments의 Bridge 역할이 설명되는가?",
    ):
        add_bullet(doc, item)

    doc.core_properties.title = "학사관리시스템 데이터베이스 설계 및 실습 요약보고서"
    doc.core_properties.subject = "PostgreSQL 학사관리시스템 ERD 및 SQL 실습"
    doc.core_properties.author = "SKALA 학습자"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()

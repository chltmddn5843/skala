from __future__ import annotations

import json
import re
import statistics
import sys
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from config import connect  # noqa: E402
from queries import QUESTIONS  # noqa: E402


REPORT_DIR = ROOT / "report"
TMP_DIR = ROOT.parent.parent / "tmp" / "practice4_report"
ASSET_DIR = TMP_DIR / "assets"
OUTPUT = REPORT_DIR / "판교_4반_최승우_종합실습4.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "203748"
MUTED = "666666"
LIGHT = "F2F4F7"
CODE_FILL = "F4F6F9"
RESULT_FILL = "E8EEF5"
INSIGHT_FILL = "F1F7E9"
CAUTION_FILL = "FFF8E8"

BEFORE_SQL = {
    1: """
        SELECT
            sum(oi.line_total) AS total_revenue
        FROM ecom.orders o
        JOIN ecom.order_items oi
          ON oi.order_id = o.order_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
          AND o.order_ts >= now() - interval '1 month'
    """,
    2: """
        SELECT
            date_trunc('month', o.order_ts)::date AS month,
            count(DISTINCT o.order_id) AS order_count,
            sum(oi.line_total) AS revenue,
            sum(oi.line_total) / count(DISTINCT o.order_id) AS aov
        FROM ecom.orders o
        JOIN ecom.order_items oi
          ON oi.order_id = o.order_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
        GROUP BY date_trunc('month', o.order_ts)::date
        ORDER BY month
    """,
    3: """
        SELECT
            c.category_id,
            c.category_name,
            sum(oi.line_total) AS revenue
        FROM ecom.order_items oi
        JOIN ecom.orders o
          ON o.order_id = oi.order_id
        JOIN ecom.products p
          ON p.product_id = oi.product_id
        JOIN ecom.categories c
          ON c.category_id = p.category_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
          AND o.order_ts >= now() - interval '90 days'
        GROUP BY
            c.category_id,
            c.category_name
        ORDER BY
            revenue DESC,
            c.category_id
        LIMIT 10
    """,
    4: """
        SELECT
            p.product_id,
            p.product_name,
            sum(oi.line_total) AS revenue,
            rank() OVER (
                ORDER BY sum(oi.line_total) DESC
            ) AS revenue_rank
        FROM ecom.products p
        JOIN ecom.order_items oi
          ON oi.product_id = p.product_id
        JOIN ecom.orders o
          ON o.order_id = oi.order_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
        GROUP BY
            p.product_id,
            p.product_name
        ORDER BY
            revenue_rank,
            p.product_id
        LIMIT 20
    """,
    5: """
        SELECT
            c.customer_id,
            c.full_name,
            current_date - max(o.order_ts)::date AS recency_days,
            count(DISTINCT o.order_id) AS frequency,
            sum(oi.line_total) AS monetary
        FROM ecom.customers c
        JOIN ecom.orders o
          ON o.customer_id = c.customer_id
        JOIN ecom.order_items oi
          ON oi.order_id = o.order_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
        GROUP BY
            c.customer_id,
            c.full_name
        ORDER BY
            monetary DESC,
            c.customer_id
    """,
    6: """
        WITH purchases AS (
            SELECT
                customer_id,
                order_ts
            FROM ecom.orders
            WHERE order_status IN ('paid', 'shipped', 'delivered')
        ),
        first_buy AS (
            SELECT
                customer_id,
                min(order_ts) AS first_order_ts
            FROM purchases
            GROUP BY customer_id
        )

        SELECT
            count(*) AS first_buyers,
            count(*) FILTER (
                WHERE EXISTS (
                    SELECT 1
                    FROM purchases p
                    WHERE p.customer_id = f.customer_id
                      AND p.order_ts > f.first_order_ts
                      AND p.order_ts <= f.first_order_ts + interval '30 days'
                )
            ) AS repurchasers,
            ecom.safe_div(
                count(*) FILTER (
                    WHERE EXISTS (
                        SELECT 1
                        FROM purchases p
                        WHERE p.customer_id = f.customer_id
                          AND p.order_ts > f.first_order_ts
                          AND p.order_ts <= f.first_order_ts + interval '30 days'
                    )
                ),
                count(*)
            ) AS repurchase_rate
        FROM first_buy f
    """,
    7: """
        SELECT
            p.product_id,
            p.product_name,
            (
                SELECT i.qty_on_hand
                FROM ecom.inventory i
                WHERE i.product_id = p.product_id
            ) AS qty_on_hand,
            (
                SELECT i.reorder_point
                FROM ecom.inventory i
                WHERE i.product_id = p.product_id
            ) AS reorder_point,
            (
                SELECT i.reorder_point - i.qty_on_hand
                FROM ecom.inventory i
                WHERE i.product_id = p.product_id
            ) AS shortage
        FROM ecom.products p
        WHERE (
            SELECT i.qty_on_hand < i.reorder_point
            FROM ecom.inventory i
            WHERE i.product_id = p.product_id
        )
        ORDER BY
            shortage DESC,
            p.product_id
    """,
    8: """
        SELECT
            p.product_id,
            p.product_name,
            avg(r.rating) AS avg_rating,
            count(*) AS review_count
        FROM ecom.products p
        JOIN ecom.reviews r
          ON r.product_id = p.product_id
        GROUP BY
            p.product_id,
            p.product_name
        HAVING avg(r.rating) >= 4.5
           AND count(*) >= 50
        ORDER BY
            avg_rating DESC,
            review_count DESC,
            p.product_id
    """,
    9: """
        SELECT
            o.coupon_code IS NOT NULL AS used_coupon,
            count(DISTINCT o.order_id) AS order_count,
            sum(oi.line_total) / count(DISTINCT o.order_id) AS avg_order_amount
        FROM ecom.orders o
        JOIN ecom.order_items oi
          ON oi.order_id = o.order_id
        WHERE o.order_status IN ('paid', 'shipped', 'delivered')
        GROUP BY o.coupon_code IS NOT NULL
        ORDER BY used_coupon
    """,
    10: """
        WITH lifetime AS (
            SELECT
                o.customer_id,
                sum(oi.line_total) AS lifetime_revenue
            FROM ecom.orders o
            JOIN ecom.order_items oi
              ON oi.order_id = o.order_id
            WHERE o.order_status IN ('paid', 'shipped', 'delivered')
            GROUP BY o.customer_id
        ),
        top_customers AS (
            SELECT
                customer_id,
                lifetime_revenue
            FROM lifetime
            ORDER BY
                lifetime_revenue DESC,
                customer_id
            LIMIT 30
        )

        SELECT
            tc.customer_id,
            tc.lifetime_revenue,
            (
                SELECT sum(oi.line_total)
                FROM ecom.orders o
                JOIN ecom.order_items oi
                  ON oi.order_id = o.order_id
                WHERE o.customer_id = tc.customer_id
                  AND o.order_status IN ('paid', 'shipped', 'delivered')
                  AND o.order_ts >= now() - interval '60 days'
            ) AS recent_60d_revenue
        FROM top_customers tc
        ORDER BY
            tc.lifetime_revenue DESC,
            tc.customer_id
    """,
}

META = {
    1: ("재무", "최근 한 달 실제 매출 규모 확인", "롤링 1개월·유효 주문 상태"),
    2: ("재무·경영", "월별 주문 수·매출·AOV 추이", "달력 월·주문 단위 평균"),
    3: ("구매·MD", "최근 90일 카테고리 성과 비교", "카테고리 매출 Top 10"),
    4: ("MD", "누적매출 상위 상품 선별", "RANK 기준 정확히 20개"),
    5: ("CRM", "고객별 최근성·빈도·금액 파악", "유효 주문·주문 단위 빈도"),
    6: ("CRM", "첫 구매 후 30일 내 재구매율 확인", "첫 주문 이후 추가 주문"),
    7: ("구매·SCM", "재주문이 필요한 상품 탐지", "현재고 < 재주문 임계치"),
    8: ("MD·마케팅", "리뷰 기반 효자상품 선별", "평점 4.5 이상·리뷰 50개 이상"),
    9: ("마케팅·재무", "쿠폰 사용 여부별 AOV 비교", "주문별 금액을 먼저 집계"),
    10: ("CRM·VIP", "상위 1% 고객의 최근 60일 매출", "누적매출 Top 30 후 60일 집계"),
}

TUNING = {
    1: ("직접 조인 집계", "기간·상태 대상 주문을 먼저 명명해 조건과 집계 경로를 분리", "PostgreSQL이 CTE를 인라인하면 계획은 같을 수 있다. 성능보다 요청 조건의 가독성이 핵심이다."),
    2: ("주문상품 행에서 COUNT DISTINCT와 합계를 동시에 수행", "주문별 금액 선집계 후 월 집계", "1:N 조인으로 부푼 행을 주문 단위로 먼저 축소했다."),
    3: ("전체 조인 후 기간 필터와 집계", "최근 판매행을 먼저 좁힌 뒤 상품·카테고리 연결", "소규모 데이터에서는 계획이 같을 수 있으나 분석 경계가 선명해졌다."),
    4: ("상품명까지 포함해 집계와 순위를 한 단계에서 처리", "상품 매출 집계→순위→상품명 조회로 역할 분리", "정확히 20개를 반환하고 동점은 product_id로 결정했다."),
    5: ("주문상품 행에서 COUNT DISTINCT로 빈도 복원", "주문별 금액을 선집계한 뒤 고객 RFM 계산", "빈도와 금액의 집계 단위를 주문으로 통일했다."),
    6: ("고객별 EXISTS 상관 서브쿼리를 반복", "윈도 함수로 첫 구매일 계산 후 BOOL_OR로 고객당 1회 집계", "반복 탐색을 단일 스캔·집계 흐름으로 바꾼 대표 개선 사례다."),
    7: ("상품마다 재고 스칼라 서브쿼리를 네 번 실행", "inventory와 products를 한 번 JOIN", "같은 테이블 반복 조회를 제거했다."),
    8: ("상품명까지 조인한 뒤 모든 리뷰를 집계", "리뷰를 상품별로 먼저 집계·필터 후 상품명 조회", "효자 후보만 마스터 데이터와 결합한다."),
    9: ("주문상품 행 합계÷COUNT DISTINCT로 AOV 계산", "주문별 금액 선집계 후 쿠폰 그룹 AVG", "AOV의 업무 단위가 주문임을 SQL 구조에 직접 반영했다."),
    10: ("Top 30 고객마다 최근 60일 매출 상관 조회", "한 번의 고객 집계에서 누적·60일 매출을 조건부 합산", "고객당 반복 스캔을 한 번의 집계로 줄였다."),
}


def set_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tcpr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tcw = cell._tc.get_or_add_tcPr().tcW
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")


def mark_header_row(table):
    trpr = table.rows[0]._tr.get_or_add_trPr()
    if trpr.find(qn("w:tblHeader")) is None:
        trpr.append(OxmlElement("w:tblHeader"))


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_font(run, size=9, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(1.0)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    set_font(header.add_run("SKALA | E-Commerce 매출 분석 및 정리"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("판교 4반 · 최승우   |   "), size=8.5, color=MUTED)
    add_page_field(footer)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text.upper()), size=10, color=BLUE, bold=True)


def add_code_block(doc, sql):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, CODE_FILL)
    cell_margins(cell, 100, 140, 100, 140)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    lines = sql.strip().splitlines()
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        set_font(paragraph.add_run(line.rstrip()), name="Menlo", size=6.8, color="1F2937")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label, text, fill=INSIGHT_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(f"{label}  "), size=9.5, color=DARK, bold=True)
    set_font(p.add_run(text), size=9.5, color="1F2937")


def add_label_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [1800, 7560])
    for label, value in rows:
        cells = table.add_row().cells
        for cell in cells:
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cells[0], LIGHT)
        set_font(cells[0].paragraphs[0].add_run(label), size=9, color=DARK, bold=True)
        set_font(cells[1].paragraphs[0].add_run(value), size=9, color="1F2937")
    set_table_geometry(table, [1800, 7560])
    return table


def add_comparison_table(doc, before, after, equal):
    rows = (
        ("Execution Time", f"{before['execution']:.3f} ms", f"{after['execution']:.3f} ms"),
        ("Planning Time", f"{before['planning']:.3f} ms", f"{after['planning']:.3f} ms"),
        ("Actual Rows", f"{before['rows']:,}", f"{after['rows']:,}"),
        ("Shared Hit", f"{before['hits']:,}", f"{after['hits']:,}"),
        ("Shared Read", f"{before['reads']:,}", f"{after['reads']:,}"),
        ("Top Plan Node", before["node"], after["node"]),
        ("결과 동일", "-", "YES" if equal else "NO"),
    )
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2880, 3240, 3240])
    for cell, value in zip(table.rows[0].cells, ("측정 항목", "튜닝 전", "튜닝 후")):
        shade(cell, RESULT_FILL)
        cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(value), size=8.8, color=DARK, bold=True)
    mark_header_row(table)
    for label, old, new in rows:
        cells = table.add_row().cells
        for cell in cells:
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_font(cells[0].paragraphs[0].add_run(label), size=8.5, color="1F2937", bold=True)
        set_font(cells[1].paragraphs[0].add_run(old), size=8.5, color="1F2937")
        set_font(cells[2].paragraphs[0].add_run(new), size=8.5, color="08783E" if label != "결과 동일" else DARK, bold=label == "결과 동일")
    set_table_geometry(table, [2880, 3240, 3240])


def measure(connection, sql, repeats=3):
    samples = []
    for _ in range(repeats):
        payload = connection.execute(
            "EXPLAIN (ANALYZE, BUFFERS, FORMAT JSON) " + sql
        ).fetchone()[0][0]
        plan = payload["Plan"]
        samples.append({
            "execution": float(payload["Execution Time"]),
            "planning": float(payload["Planning Time"]),
            "rows": int(plan["Actual Rows"]),
            "hits": int(plan.get("Shared Hit Blocks", 0)),
            "reads": int(plan.get("Shared Read Blocks", 0)),
            "node": plan["Node Type"],
        })
    return {
        "execution": statistics.median(item["execution"] for item in samples),
        "planning": statistics.median(item["planning"] for item in samples),
        "rows": int(statistics.median(item["rows"] for item in samples)),
        "hits": int(statistics.median(item["hits"] for item in samples)),
        "reads": int(statistics.median(item["reads"] for item in samples)),
        "node": samples[-1]["node"],
    }


def fetch(connection, sql):
    cursor = connection.execute(sql)
    return [column.name for column in cursor.description], cursor.fetchall()


def display(value):
    if value is None:
        return "NULL"
    if isinstance(value, Decimal):
        return f"{value:,.2f}"
    return str(value)


def make_result_image(path, title, columns, rows, elapsed, max_rows=8):
    font_path = "/System/Library/Fonts/AppleSDGothicNeo.ttc"
    mono_path = "/System/Library/Fonts/Menlo.ttc"
    title_font = ImageFont.truetype(font_path, 30)
    mono = ImageFont.truetype(mono_path, 22)
    shown = rows[:max_rows]
    values = [[display(value) for value in row] for row in shown]
    widths = [
        min(25, max([len(name), *(len(row[index]) for row in values)]))
        for index, name in enumerate(columns)
    ]

    def line(items):
        return " | ".join(str(value)[:width].ljust(width) for value, width in zip(items, widths))

    lines = [line(columns), "-+-".join("-" * width for width in widths)]
    lines.extend(line(row) for row in values)
    if len(rows) > max_rows:
        lines.append(f"... {len(rows) - max_rows:,} rows omitted")
    lines.append(f"rows={len(rows):,} | median execution={elapsed:.3f} ms")
    height = 130 + len(lines) * 34
    image = Image.new("RGB", (1600, height), "#111827")
    draw = ImageDraw.Draw(image)
    draw.text((38, 28), title, font=title_font, fill="#F9FAFB")
    draw.multiline_text((38, 92), "\n".join(lines), font=mono, fill="#E5E7EB", spacing=10)
    image.save(path)


def summary(number, rows):
    if number == 1:
        return f"최근 한 달 실제 매출은 {display(rows[0][0])}이다."
    if number == 2:
        month, orders, revenue, aov = rows[-1]
        return f"최신 월({month}) 주문은 {orders:,}건, 매출 {display(revenue)}, AOV {display(aov)}로 집계됐다. 부분 월 여부를 함께 전달해야 한다."
    if number == 3:
        _, name, revenue = rows[0]
        return f"최근 90일 1위 카테고리는 {name}, 매출은 {display(revenue)}다. 순위는 랜덤 시드 재적재 시 바뀔 수 있다."
    if number == 4:
        product_id, name, revenue, rank = rows[0]
        return f"누적매출 1위는 {name}(ID {product_id}), 매출 {display(revenue)}다. 결과는 정확히 20개로 고정했다."
    if number == 5:
        customer_id, name, recency, frequency, monetary = rows[0]
        return f"최고 Monetary 고객은 {name}(ID {customer_id})이며 구매 {frequency}회, 누적금액 {display(monetary)}, 최근성 {recency}일이다."
    if number == 6:
        first_buyers, repurchasers, rate = rows[0]
        return f"첫 구매 고객 {first_buyers:,}명 중 {repurchasers:,}명이 30일 내 재구매해 재구매율은 {float(rate):.1%}다."
    if number == 7:
        product_id, name, qty, point, shortage = rows[0]
        return f"재고 임계치 미달 상품은 {len(rows):,}개다. 가장 긴급한 {name}(ID {product_id})은 현재고 {qty}, 임계치 {point}, 부족분 {shortage}다."
    if number == 8:
        product_id, name, rating, reviews = rows[0]
        return f"효자상품 조건을 충족한 상품은 {len(rows):,}개다. 상위 {name}(ID {product_id})은 평점 {float(rating):.2f}, 리뷰 {reviews}개다."
    if number == 9:
        mapping = {bool(row[0]): row for row in rows}
        no_coupon = mapping[False][2]
        coupon = mapping[True][2]
        lift = float(coupon / no_coupon - 1)
        return f"쿠폰 사용 주문 AOV는 {display(coupon)}, 미사용 주문은 {display(no_coupon)}로 표본상 {lift:.1%} 높다. 인과효과가 아니라 관찰 비교다."
    customer_id, lifetime, recent = rows[0]
    total_recent = sum((row[2] or Decimal(0)) for row in rows)
    return f"상위 1%는 30명으로 고정했다. 1위 고객 ID {customer_id}의 누적매출은 {display(lifetime)}, 최근 60일 매출은 {display(recent)}이며 Top 30의 최근 매출 합계는 {display(total_recent)}다."


def collect():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    data = {}
    with connect() as connection:
        database, user, version, superuser = connection.execute(
            """
            SELECT
                current_database(),
                current_user,
                version(),
                rolsuper
            FROM pg_roles
            WHERE rolname = current_user
            """
        ).fetchone()
        counts = dict(connection.execute("""
            SELECT 'customers', count(*) FROM ecom.customers
            UNION ALL SELECT 'products', count(*) FROM ecom.products
            UNION ALL SELECT 'orders', count(*) FROM ecom.orders
            UNION ALL SELECT 'order_items', count(*) FROM ecom.order_items
            UNION ALL SELECT 'reviews', count(*) FROM ecom.reviews
        """).fetchall())

        for number in range(1, 11):
            before_sql = QUESTIONS[number]["before_sql"].strip()
            after_sql = QUESTIONS[number]["after_sql"].strip()
            before_columns, before_rows = fetch(connection, before_sql)
            after_columns, after_rows = fetch(connection, after_sql)
            assert before_rows == after_rows, f"Q{number:02d} 튜닝 전후 결과 불일치"
            before_metrics = measure(connection, before_sql)
            after_metrics = measure(connection, after_sql)
            before_image = ASSET_DIR / f"Q{number:02d}_before.png"
            after_image = ASSET_DIR / f"Q{number:02d}_after.png"
            make_result_image(before_image, f"Q{number:02d} 튜닝 전 결과", before_columns, before_rows, before_metrics["execution"])
            make_result_image(after_image, f"Q{number:02d} 튜닝 후 결과", after_columns, after_rows, after_metrics["execution"])
            data[number] = {
                "before_sql": before_sql,
                "after_sql": after_sql,
                "before_rows": before_rows,
                "after_rows": after_rows,
                "before_metrics": before_metrics,
                "after_metrics": after_metrics,
                "before_image": before_image,
                "after_image": after_image,
                "equal": True,
                "summary": summary(number, after_rows),
            }
    return data, {
        "database": database,
        "user": user,
        "version": version,
        "superuser": superuser,
        "counts": counts,
    }


def add_picture(doc, path, width=6.15):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    doc_pr = doc.inline_shapes[-1]._inline.docPr
    doc_pr.set("descr", path.stem.replace("_", " "))
    doc_pr.set("title", path.stem)


def add_plan_log(doc, path, label, max_height=1800):
    """logs의 원본 PNG를 읽기 좋은 높이로만 나눠 그대로 첨부한다."""
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as source:
        parts = []
        for index, top in enumerate(range(0, source.height, max_height), start=1):
            if source.height <= max_height:
                parts = [path]
                break
            part = ASSET_DIR / f"{path.stem}_part{index}.png"
            source.crop((0, top, source.width, min(top + max_height, source.height))).save(part)
            parts.append(part)

    for index, part in enumerate(parts, start=1):
        suffix = f" ({index}/{len(parts)})" if len(parts) > 1 else ""
        heading = doc.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(3)
        set_font(heading.add_run(label + suffix), size=9.5, color=DARK, bold=True)
        add_picture(doc, part)


def add_delivery_block(doc, team, judgment, message):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    row = table.rows[0]
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))
    cell = row.cells[0]
    shade(cell, RESULT_FILL)
    cell_margins(cell, 160, 180, 160, 180)
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(6)
    set_font(first.add_run("튜닝 판단\n"), size=9.5, color=DARK, bold=True)
    set_font(first.add_run(judgment), size=9.5, color="1F2937")
    second = cell.add_paragraph()
    second.paragraph_format.space_after = Pt(0)
    set_font(second.add_run(f"{team} 전달안\n"), size=9.5, color=DARK, bold=True)
    set_font(second.add_run(message), size=9.5, color="1F2937")


def plan_log_stats(filename, strategy=None):
    text = (ROOT / "logs" / filename).read_text(encoding="utf-8")
    plan = text.split("\nPLAN\n", 1)[1]
    lines = plan.splitlines()
    line = next(
        item for item in lines
        if "actual time=" in item and (strategy is None or strategy in item)
    )
    node = strategy or line.split("  (cost", 1)[0].strip()
    inner_scan = next(
        (item for item in lines if "Index Scan using idx_order_items_order" in item),
        line,
    )
    return {
        "node": node,
        "estimated_rows": int(re.search(r"\brows=(\d+)", line).group(1)),
        "actual_rows": int(re.search(r"actual time=.*?\brows=(\d+)", line).group(1)),
        "loops": int(re.search(r"\bloops=(\d+)", line).group(1)),
        "inner_loops": int(re.search(r"\bloops=(\d+)", inner_scan).group(1)),
        "hits": int(re.search(r"Buffers: shared hit=(\d+)", plan).group(1)),
        "execution": float(re.search(r"Execution Time: ([\d.]+) ms", plan).group(1)),
        "memory": int(re.findall(r"Memory Usage: (\d+)kB", plan)[-1]) if "Memory Usage:" in plan else 0,
    }


def detected_join(filename):
    text = (ROOT / "logs" / filename).read_text(encoding="utf-8")
    return next(name for name in ("Nested Loop", "Hash Join", "Merge Join") if name in text)


def add_question(doc, number, item):
    title = QUESTIONS[number]["title"]
    team, purpose, criteria = META[number]
    doc.add_page_break()
    doc.add_heading(f"Q{number}. {title}", level=1)
    add_label_table(doc, (("요청 부서", team), ("요청 목적", purpose), ("판정 기준", criteria)))

    before_problem, after_change, judgment = TUNING[number]
    doc.add_heading("튜닝 전", level=2)
    doc.add_paragraph(before_problem)
    add_code_block(doc, item["before_sql"])
    doc.add_heading("튜닝 전 실행 결과", level=3)
    add_picture(doc, item["before_image"])
    add_plan_log(
        doc,
        ROOT / "logs" / f"Q{number:02d}_plan_before.png",
        f"튜닝 전 성능 원문 로그 · Q{number:02d}_plan_before.png",
    )

    doc.add_heading("튜닝 포인트", level=2)
    add_callout(doc, "변경", after_change, CAUTION_FILL)

    doc.add_heading("튜닝 후", level=2)
    add_code_block(doc, item["after_sql"])
    doc.add_heading("튜닝 후 실행 결과", level=3)
    add_picture(doc, item["after_image"])
    add_plan_log(
        doc,
        ROOT / "logs" / f"Q{number:02d}_plan_after.png",
        f"튜닝 후 성능 원문 로그 · Q{number:02d}_plan_after.png",
    )

    doc.add_heading("튜닝 전후 실행계획 비교", level=2)
    add_comparison_table(doc, item["before_metrics"], item["after_metrics"], item["equal"])
    doc.add_heading("업무 판단 및 요청 부서 전달안", level=2)
    add_delivery_block(doc, team, judgment, item["summary"])


def add_q11_section(doc):
    doc.add_page_break()
    doc.add_heading("Q11. 안전한 나눗셈 함수 비교", level=1)
    add_label_table(doc, (
        ("검증 목적", "분모 0·NULL 처리 차이 확인"),
        ("비교 함수", "ecom.f_safe_div · ecom.safe_div"),
        ("적용 판단", "평균 계산은 관측 없음과 실제 0을 구분"),
    ))
    doc.add_heading("검증 SQL 및 결과", level=2)
    add_code_block(doc, QUESTIONS[11]["sql"])
    image = ROOT / "logs" / "Q11.png"
    if image.exists():
        add_picture(doc, image, 6.15)
    add_callout(doc, "결과", "정상 분모에서는 두 함수 모두 5를 반환했다. 분모가 0이면 f_safe_div는 0, safe_div는 NULL을 반환하며 safe_div는 NULL 분모도 NULL로 유지한다.", RESULT_FILL)
    add_callout(doc, "선택", "평균값에서 ‘관측 없음’과 실제 0을 구분하기 위해 Q2·Q6에는 ecom.safe_div를 사용했다.")


def add_join_section(doc):
    nested = plan_log_stats("join_nested_loop.log", "Nested Loop")
    hashed = plan_log_stats("join_hash_join.log", "Hash Join")
    merged = plan_log_stats("join_merge_join.log", "Merge Join")
    default_strategy = detected_join("join_default.log")
    doc.add_page_break()
    doc.add_heading("부록 A. 조인 방식 비교", level=1)
    doc.add_paragraph("같은 고객별 매출 집계 쿼리에서 PostgreSQL의 조인 옵션을 각각 강제해 실제 실행계획을 비교했다. 목적은 특정 조인을 암기하는 것이 아니라, 입력 행 수·선택도·인덱스·반복 탐색 비용을 보고 옵티마이저의 선택을 검증하는 것이다. 수치는 현재 시드 데이터와 캐시 상태에서의 1회 관측값이다.")
    doc.add_heading("이 쿼리에서 조인 전략을 고민한 이유", level=2)
    for text in (
        "orders 9,890건 중 유효 상태 7,613건이 남아 외부 입력이 작지 않고, order_items 26,716건과 1:N으로 결합된다.",
        "조인 조건은 order_id 동등 비교이므로 Hash Join의 기본 조건에 맞지만, 양쪽에 order_id 인덱스가 있어 Merge Join과 Nested Loop도 후보가 된다.",
        "최종 집계는 2,250명이지만 조인 단계에서는 20,557행을 처리하므로, 출력 행 수만 보고 Nested Loop가 적합하다고 판단하면 안 된다.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(text), size=9.5, color="1F2937")

    doc.add_heading("확인 절차와 판독 포인트", level=2)
    add_label_table(doc, (
        ("1. 비교 통제", "동일 SQL·동일 데이터에서 SET LOCAL로 하나의 조인 방식만 활성화"),
        ("2. 결과 동일성", "세 계획 모두 조인 20,557행, 최종 2,250행인지 확인"),
        ("3. 추정 품질", "예상 20,565행 대비 실제 20,557행으로 행 수 오차가 매우 작음"),
        ("4. 실제 비용", "Execution Time·Buffers hit·loops·스캔 방식·Hash Batch/메모리를 함께 비교"),
        ("5. 해석 제한", "조인 강제는 원인 확인용이며, 운영 SQL에서는 통계정보와 옵티마이저 선택을 우선"),
    ))
    add_callout(doc, "기본 선택", f"모든 조인 옵션을 활성화한 기본 실행계획에서도 PostgreSQL은 {default_strategy}을 선택했다. 강제 비교 결과와 기본 선택이 일치하는지 먼저 확인했다.", RESULT_FILL)
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1600, 1660, 1880, 1880, 2340])
    headers = ("방식", "실행시간", "Buffers hit", "적합한 상황", "이번 결과")
    for cell, value in zip(table.rows[0].cells, headers):
        shade(cell, RESULT_FILL)
        cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(value), size=8.2, color=DARK, bold=True)
    mark_header_row(table)
    rows = (
        ("Nested Loop", f"{nested['execution']:.3f} ms", f"{nested['hits']:,}", "외부 결과가 작고 내부 인덱스가 유효", f"내부 Index Scan loops={nested['inner_loops']:,}"),
        ("Hash Join", f"{hashed['execution']:.3f} ms", f"{hashed['hits']:,}", "큰 동등 조인·정렬 불필요", f"Batches=1, Hash Memory={hashed['memory']:,}kB"),
        ("Merge Join", f"{merged['execution']:.3f} ms", f"{merged['hits']:,}", "양쪽 입력이 조인 키로 정렬", "양쪽 인덱스 순서를 활용"),
    )
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cell.paragraphs[0].add_run(value), size=7.8, color="1F2937")
    set_table_geometry(table, [1600, 1660, 1880, 1880, 2340])
    doc.add_heading("실측 결과 해석", level=2)
    add_callout(doc, "Nested Loop", f"orders 7,613행마다 idx_order_items_order를 탐색해 내부 Index Scan loops={nested['inner_loops']:,}이 발생했다. 누적 Buffers hit={nested['hits']:,}, {nested['execution']:.3f}ms로 가장 비싼 계획이었다.", CAUTION_FILL)
    add_callout(doc, "Hash Join", f"필터를 통과한 orders로 {hashed['memory']:,}kB 해시를 만들고 order_items를 1회 순차 스캔했다. 디스크 spill 없이 Buffers hit={hashed['hits']:,}, {hashed['execution']:.3f}ms를 기록했다.", RESULT_FILL)
    add_callout(doc, "Merge Join", f"orders_pkey와 idx_order_items_order의 정렬 순서를 활용해 별도 Sort 없이 실행됐다. Buffers hit={merged['hits']:,}, {merged['execution']:.3f}ms로 Hash Join과 비교했다.")
    speed_gain = nested["execution"] - hashed["execution"]
    speed_pct = speed_gain / nested["execution"]
    buffer_pct = 1 - hashed["hits"] / nested["hits"]
    add_callout(doc, "결론", f"현 분포에서 Hash Join은 Nested Loop보다 {speed_gain:.3f}ms({speed_pct:.1%}) 빨랐고 버퍼 hit를 {buffer_pct:.1%} 줄였다. 기간 필터로 외부 행이 소수로 줄면 Nested Loop를, 양쪽이 이미 정렬된 대량 입력이면 Merge Join을 다시 검토한다.", INSIGHT_FILL)

    doc.add_heading("문항별 조인 전략 고민 포인트", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [900, 3200, 5260])
    for cell, value in zip(table.rows[0].cells, ("문항", "실제 계획에서 확인한 점", "판단")):
        shade(cell, RESULT_FILL)
        cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(value), size=8.4, color=DARK, bold=True)
    mark_header_row(table)
    for values in (
        ("Q3", "orders→order_items→products→categories에서 Hash Join 연속 사용", "90일 조건도 주문 5,873건을 남겨 입력이 크므로 Hash가 합리적. 기간 선택도가 크게 낮아지면 조인 순서와 인덱스를 재확인"),
        ("Q5", "주문상품 20,557행 선집계 뒤 customers와 Hash Join", "조인 종류보다 주문 단위 선집계가 먼저. 1:N 중복을 줄인 뒤 고객 마스터를 결합"),
        ("Q6", "튜닝 후 다중 테이블 조인 없이 orders 1회 Scan+Window", "조인 방식 강제보다 반복 EXISTS 제거가 핵심. 상관 탐색을 단일 집계 흐름으로 변경"),
        ("Q7", "products 600행과 부족재고 63행을 Hash Join", "작은 테이블에서는 Seq Scan+Hash가 스칼라 서브쿼리 반복보다 단순하고 저렴"),
        ("Q10", "유효 주문상품 20,557행을 Hash Join 후 고객 집계", "Top 30마다 재조회하지 않고 전체 판매행을 한 번 결합·조건부 집계"),
    ):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cell.paragraphs[0].add_run(value), size=7.8, color="1F2937")
    set_table_geometry(table, [900, 3200, 5260])

    doc.add_heading("Bitmap Heap Scan은 어떻게 볼 것인가", level=2)
    doc.add_paragraph("Bitmap Heap Scan은 Hash/Nested Loop/Merge와 같은 조인 알고리즘이 아니라 테이블 접근 방식이다. 인덱스 조건에 맞는 행 위치를 비트맵으로 모은 뒤 필요한 heap 페이지를 묶어 읽으므로, 한 건 조회와 전체 조회 사이의 중간 선택도에서 유리할 수 있다.")
    index_scan = plan_log_stats("access_index_scan.log")
    bitmap_scan = plan_log_stats("access_bitmap_heap_scan.log")
    seq_scan = plan_log_stats("access_seq_scan.log")
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [2300, 1600, 1900, 1500, 2060])
    for cell, value in zip(table.rows[0].cells, ("확인 조건", "예상/실제 행", "선택된 접근 방식", "Buffers hit", "Execution Time")):
        shade(cell, RESULT_FILL)
        cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(value), size=8.1, color=DARK, bold=True)
    mark_header_row(table)
    for condition, stats in (
        ("PK order_id=1", index_scan),
        ("customer_id 1~100", bitmap_scan),
        ("customer_id 1~3000", seq_scan),
    ):
        cells = table.add_row().cells
        values = (
            condition,
            f"{stats['estimated_rows']:,}/{stats['actual_rows']:,}",
            stats["node"],
            f"{stats['hits']:,}",
            f"{stats['execution']:.3f} ms",
        )
        for cell, value in zip(cells, values):
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_font(cell.paragraphs[0].add_run(value), size=7.8, color="1F2937")
    set_table_geometry(table, [2300, 1600, 1900, 1500, 2060])
    add_callout(doc, "Bitmap 실측", f"customer_id 1~100 조건은 예상 {bitmap_scan['estimated_rows']:,}행, 실제 {bitmap_scan['actual_rows']:,}행으로 추정이 정확했다. idx_orders_customer_ts가 만든 비트맵으로 heap 15개 블록을 묶어 읽어 {bitmap_scan['execution']:.3f}ms를 기록했다. 전체 범위에서는 인덱스 왕복보다 Seq Scan이 선택됐다.", RESULT_FILL)
    add_callout(doc, "판단 기준", "접근 방식은 서로 다른 결과 건수의 속도 경쟁이 아니다. 예상/실제 행 수, 선택도, Heap Blocks, Recheck Cond, Buffers를 보고 옵티마이저가 구간별로 Index→Bitmap→Seq Scan을 전환하는지 확인한다.")

    doc.add_heading("접근 방식 실행계획 증적", level=2)
    for filename, caption in (
        ("access_index_scan.png", "정확한 PK 1건: Index Scan"),
        ("access_bitmap_heap_scan.png", "중간 선택도 992건: Bitmap Index Scan + Bitmap Heap Scan"),
        ("access_seq_scan.png", "전체 범위 9,890건: Seq Scan"),
    ):
        add_plan_log(doc, ROOT / "logs" / filename, caption)

    doc.add_heading("실행계획 증적", level=2)
    for name, caption in (
        ("join_default.png", f"기본 옵티마이저 선택: {default_strategy}"),
        ("join_nested_loop.png", "Nested Loop: 내부 Index Scan loops=7,613과 Buffers hit=23,069 확인"),
        ("join_hash_join.png", "Hash Join: Hash Batches=1, Memory Usage=421kB, Buffers hit=360 확인"),
        ("join_merge_join.png", "Merge Join: 양쪽 Index Scan과 별도 Sort 없음 확인"),
    ):
        image = ROOT / "logs" / name
        if image.exists():
            p = doc.add_paragraph()
            set_font(p.add_run(caption), size=8.5, color=MUTED, bold=True)
            add_picture(doc, image, 6.15)

    doc.add_heading("비교 SQL", level=2)
    add_code_block(doc, """SELECT
    o.customer_id,
    count(*) AS item_count,
    sum(oi.line_total) AS revenue
FROM ecom.orders o
JOIN ecom.order_items oi
  ON oi.order_id = o.order_id
WHERE o.order_status IN ('paid', 'shipped', 'delivered')
GROUP BY o.customer_id;""")


def add_mv_section(doc):
    doc.add_page_break()
    doc.add_heading("부록 B. Materialized View", level=1)
    doc.add_paragraph("반복 요청이 많은 일별 GMV 집계를 미리 저장해 조회 비용을 줄이는 구조다. 일반 View와 달리 결과를 물리적으로 보관하므로 원본 변경 후 명시적으로 갱신해야 한다.")
    doc.add_heading("생성 및 갱신 스크립트", level=2)
    add_code_block(doc, """CREATE MATERIALIZED VIEW IF NOT EXISTS ecom.mv_daily_gmv AS
SELECT
    date_trunc('day', o.order_ts) AS day,
    sum(oi.line_total) AS gmv
FROM ecom.orders o
JOIN ecom.order_items oi
  ON oi.order_id = o.order_id
WHERE o.order_status IN ('paid', 'shipped', 'delivered')
GROUP BY date_trunc('day', o.order_ts);

CREATE UNIQUE INDEX IF NOT EXISTS ux_mv_daily_gmv_day
ON ecom.mv_daily_gmv(day);

REFRESH MATERIALIZED VIEW ecom.mv_daily_gmv;

SELECT
    count(*) AS days,
    sum(gmv) AS gmv,
    max(day) AS latest_day
FROM ecom.mv_daily_gmv;""")
    image = ROOT / "logs" / "materialized_view.png"
    if image.exists():
        add_picture(doc, image, 6.15)
    add_callout(doc, "검증 결과", "트랜잭션 안에서 테스트 매출 123.45를 추가하자 원본 GMV만 증가하고 MV는 이전 값을 유지했다. REFRESH 후 두 합계가 같아졌으며, 마지막 ROLLBACK으로 테스트 행이 남지 않았음을 확인했다.", RESULT_FILL)
    add_callout(doc, "운영 판단", "Materialized View는 자동 갱신되지 않는다. day UNIQUE 인덱스를 생성했으므로 운영 조회 중단을 줄이려면 별도 트랜잭션에서 REFRESH MATERIALIZED VIEW CONCURRENTLY를 사용할 수 있다.", CAUTION_FILL)

    doc.add_heading("오후 3시 기준 갱신 운영안", level=2)
    doc.add_paragraph("다른 팀에 전달하는 일별 매출 자료는 Asia/Seoul 기준 오후 3시 스냅샷으로 정의한다. ‘오후 3시 기준’은 주문 시각 필터가 아니라, 15:00까지 원본에 커밋된 데이터를 MV에 반영하는 갱신 시점이다.")
    add_label_table(doc, (
        ("14:55", "당일 적재 완료·지연 작업·잠금 여부 확인"),
        ("15:00", "REFRESH MATERIALIZED VIEW CONCURRENTLY ecom.mv_daily_gmv 실행"),
        ("15:05", "max(day), sum(gmv), 원본 GMV와 차이 검증 후 구매·재무팀 전달"),
        ("지연 데이터", "15시 이후 보정 적재가 있으면 완료 이벤트 뒤 재갱신하고 전달본에 재산출 시각 표시"),
        ("장애 대응", "갱신 실패 시 이전 스냅샷을 최신값으로 오인하지 않도록 기준시각과 실패 상태를 함께 공지"),
    ))
    add_code_block(doc, """REFRESH MATERIALIZED VIEW CONCURRENTLY ecom.mv_daily_gmv;

SELECT
    max(day) AS latest_day,
    sum(gmv) AS gmv
FROM ecom.mv_daily_gmv;""")


def add_appendix(doc):
    doc.add_page_break()
    doc.add_heading("부록 C. 최종 인사이트", level=1)
    doc.add_heading("랜덤 데이터에서 배운 점", level=2)
    insights = (
        "숫자 자체보다 정의가 먼저다. 매출 상태·기간·주문 단위를 합의하지 않으면 같은 요청에도 다른 답이 나온다.",
        "1:N 조인은 주문 수와 AOV를 쉽게 부풀린다. 주문별 선집계는 성능보다 먼저 정확성을 지키는 장치다.",
        "튜닝 후 SQL이 항상 더 빠르지는 않았다. 작은 테이블에서는 Seq Scan이나 CTE 인라인이 합리적이며, 실행계획과 버퍼를 함께 봐야 한다.",
        "랜덤 시드는 결과값을 바꾸지만 검증 가능한 로직은 유지된다. 따라서 전달 자료에는 기준시각·조건·단위·가정을 함께 적어야 한다.",
        "상관관계와 인과관계를 구분해야 한다. 쿠폰 주문의 높은 AOV는 쿠폰 효과일 수도 있지만 고가 상품 편향이나 시드 생성 규칙의 영향일 수도 있다.",
    )
    for insight in insights:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        set_font(p.add_run(insight), size=10, color="1F2937")

    doc.add_heading("DBMS별 실행계획 확인 도구(선택)", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 2900, 4260])
    for cell, value in zip(table.rows[0].cells, ("DBMS", "대표 도구", "확인 포인트")):
        shade(cell, RESULT_FILL)
        cell_margins(cell)
        set_font(cell.paragraphs[0].add_run(value), size=8.5, color=DARK, bold=True)
    mark_header_row(table)
    for values in (
        ("PostgreSQL", "EXPLAIN (ANALYZE, BUFFERS)", "실제 행·시간·버퍼와 예상치 비교"),
        ("MySQL", "EXPLAIN ANALYZE", "실제 반복·행 수와 비용 확인"),
        ("Oracle", "DBMS_XPLAN", "실행 통계와 접근 경로 확인"),
        ("SQL Server", "Actual Execution Plan", "연산자 비용·실제/예상 행 차이 확인"),
    ):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell_margins(cell)
            set_font(cell.paragraphs[0].add_run(value), size=8.3, color="1F2937")
    set_table_geometry(table, [2200, 2900, 4260])


def build():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    data, environment = collect()
    doc = Document()
    configure(doc)

    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    add_kicker(doc, "SKALA Full-stack Engineering · 종합실습 4")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("E-Commerce 매출 분석 및 정리"), size=27, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("랜덤 데이터 위에서 요청을 분석 조건으로 바꾸고, 검증 가능한 결과로 전달하기"), size=13, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("판교 4반 · 최승우\nPostgreSQL 17.10 · ecom schema\n2026-08-14"), size=11, color=DARK, bold=True)
    add_callout(doc, "보고서 관점", "이번 실습은 가장 빠른 SQL을 고르는 시험이 아니라, 구매·재무·CRM의 요청을 같은 기준으로 재현하고 설명하는 연습이다.", RESULT_FILL)

    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("분석의 핵심은 ‘누가 요청했는가’보다 ‘무엇을 한 건으로 볼 것인가’를 먼저 정하는 것이었다. 주문과 주문상품의 1:N 관계를 그대로 집계하면 주문 수와 평균 주문 금액이 왜곡되므로, Q2·Q5·Q9는 주문 단위 선집계를 기준으로 재작성했다. 상관 서브쿼리가 반복되던 Q6·Q7·Q10은 윈도 함수·JOIN·조건부 집계로 단순화했다.")
    add_label_table(doc, (
        ("분석 DB", f"{environment['database']} / user={environment['user']}"),
        ("데이터 규모", f"고객 {environment['counts']['customers']:,} · 상품 {environment['counts']['products']:,} · 주문 {environment['counts']['orders']:,} · 주문상품 {environment['counts']['order_items']:,}"),
        ("유효 매출 상태", "paid · shipped · delivered"),
        ("공통 매출식", "order_items.line_total = unit_price × qty - discount"),
        ("측정 방식", "비교표는 EXPLAIN 3회 중앙값 · 원문 로그 화면은 실행 당시 1회 실제 출력"),
    ))
    if environment["superuser"]:
        add_callout(doc, "권한 주의", "본 실습은 로컬 관리자 계정으로 수행했다. 운영 환경에서는 분석용 읽기 전용 계정과 Materialized View 유지보수 계정을 분리하고 최소 권한만 부여해야 한다.", CAUTION_FILL)
    doc.add_heading("요청 대응 원칙", level=2)
    for text in (
        "기간·상태·집계 단위·출력 건수를 SQL보다 먼저 확정한다.",
        "랜덤 데이터에서는 수치보다 결과가 바뀌어도 유지되는 판정 로직을 검증한다.",
        "실행시간만 보지 않고 실제 행 수와 Buffers를 함께 비교한다.",
        "결과에는 기준시각과 가정을 붙이고, 관찰 비교를 인과효과로 표현하지 않는다.",
    ):
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(text), size=10, color="1F2937")

    for number in range(1, 11):
        add_question(doc, number, data[number])
    add_q11_section(doc)
    add_join_section(doc)
    add_mv_section(doc)
    add_appendix(doc)

    properties = doc.core_properties
    properties.title = "판교_4반_최승우_종합실습4"
    properties.subject = "E-Commerce 매출 분석 및 SQL 튜닝 결과 보고서"
    properties.author = "최승우"
    properties.keywords = "PostgreSQL, E-Commerce, SQL, 튜닝, SKALA"
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "questions": 11}, ensure_ascii=False))


if __name__ == "__main__":
    build()
